from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import io
from .. import models

def generate_pdf(question_set: models.QuestionSet, include_answers: bool = False) -> io.BytesIO:
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50
    
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, y, question_set.title)
    y -= 30
    
    p.setFont("Helvetica", 12)
    
    for i, q in enumerate(question_set.questions, 1):
        if y < 50:
            p.showPage()
            y = height - 50
            p.setFont("Helvetica", 12)
            
        p.drawString(50, y, f"{i}. {q.question_text}")
        y -= 20
        
        if q.question_type == "MCQ":
            import json
            options = json.loads(q.options)
            for opt in options:
                p.drawString(70, y, f"- {opt}")
                y -= 15
        
        if include_answers:
            p.setFont("Helvetica-Oblique", 10)
            p.drawString(70, y, f"Answer: {q.correct_answer}")
            p.setFont("Helvetica", 12)
            y -= 20
            
        y -= 10
        
    p.save()
    buffer.seek(0)
    return buffer

def generate_docx(question_set: models.QuestionSet, include_answers: bool = False) -> io.BytesIO:
    doc = Document()
    doc.add_heading(question_set.title, 0)
    
    for i, q in enumerate(question_set.questions, 1):
        doc.add_paragraph(f"{i}. {q.question_text}", style='List Number')
        
        if q.question_type == "MCQ":
            import json
            options = json.loads(q.options)
            for opt in options:
                doc.add_paragraph(opt, style='List Bullet')
        
        if include_answers:
            doc.add_paragraph(f"Answer: {q.correct_answer}", style='Intense Quote')
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_text(question_set: models.QuestionSet, include_answers: bool = False) -> str:
    text = f"{question_set.title}\n\n"
    for i, q in enumerate(question_set.questions, 1):
        text += f"{i}. {q.question_text}\n"
        if q.question_type == "MCQ":
            import json
            options = json.loads(q.options)
            for opt in options:
                text += f"  - {opt}\n"
        if include_answers:
            text += f"  Answer: {q.correct_answer}\n"
        text += "\n"
    return text
